from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Define a function to populate a business case template
def populate_business_case(template_path, output_path, content_dict):
    # Load the template document
    doc = Document(template_path)

    for section, content in content_dict.items():
        # Find the section header in the template
        for paragraph in doc.paragraphs:
            if section in paragraph.text:
                # Move to the next paragraph to insert content
                paragraph = paragraph._element.getparent().index(paragraph._element) + 1
                doc.paragraphs[paragraph].alignment = WD_ALIGN_PARAGRAPH.LEFT  # Align content to left
                doc.paragraphs[paragraph].add_run(content)

    # Save the populated document
    doc.save(output_path)

if __name__ == "__main__":
    # Define the paths for the template and the output document
    template_path = "business_case_template.docx"  # Replace with your template path
    output_path = "populated_business_case.docx"  # Output business case document

    # Define the content for each section (replace with your content)
    content_dict = {
        "Executive Summary": "This is the executive summary content.",
        "Business Opportunity": "This is the business opportunity content.",
        "Project Scope": "This is the project scope content.",
        # Add content for other sections as needed
    }

    # Populate the template with content and save the output document
    populate_business_case(template_path, output_path, content_dict)
    
    print(f"Populated business case document saved as '{output_path}'.")
